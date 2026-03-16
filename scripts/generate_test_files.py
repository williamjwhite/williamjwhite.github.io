import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# ---------------------------------------------------------
# CONTENT (Your real service description + industries)
# ---------------------------------------------------------

CONTENT = """
Client Document Portal — Test File
This document is auto‑generated for testing purposes only.

------------------------------------------------------------
SERVICES & EXPERTISE
------------------------------------------------------------
I design and automate secure, scalable DocuSign workflows and API integrations
for enterprise organizations and high‑volume operational teams.

As a former Partner TCSM at DocuSign, I reviewed and advised on new partner
integrations including API design, workflow architecture, backup and exception
handling, testing, go‑live readiness, user adoption, and ensuring each deployment
delivered measurable ROI.

My background includes supporting large‑scale DocuSign implementations used by
Fortune‑level brands and nationwide dealer networks across automotive,
motorcycle, and heavy equipment industries, as well as national retail, banking,
and insurance organizations.

I also develop workflows and integrations for:
Adobe Acrobat Sign, Dropbox Sign, PandaDoc, OneSpan, and other eSignature platforms.

------------------------------------------------------------
INDUSTRIES SERVED
------------------------------------------------------------
Enterprise & Regulated:
- Fortune 100/500 organizations
- Large enterprise and distributed operations
- Fintech platforms
- Banking & financial services
- Digital mortgage & eVaulting
- Capital markets, chattel, and securitization
- Compliance‑driven and regulated environments

Small & Mid‑Sized Businesses:
- Real estate teams and brokerages
- Tax offices and accounting firms
- Medical practices and clinics
- Auto, powersports, and RV dealerships
- Retail, service, and multi‑location businesses
- SMEs and growing organizations

------------------------------------------------------------
BRAND PALETTE
------------------------------------------------------------
#57C4DC  #70CDE1  #0A0A0A  #FFFFFF
#F5F5F5  #F3F4F6  #6B7280  #E5E7EB
#23272F  #2A2F38  #FAFAFA

------------------------------------------------------------
TEST FILE WATERMARK
------------------------------------------------------------
[TEST FILE — NOT FOR PRODUCTION]

------------------------------------------------------------
FOOTER
------------------------------------------------------------
Generated for testing purposes only — williamjwhite.me
"""

# ---------------------------------------------------------
# Ensure output directory exists
# ---------------------------------------------------------
os.makedirs("test-files", exist_ok=True)

# ---------------------------------------------------------
# 1. TXT File
# ---------------------------------------------------------
with open("test-files/test-file.txt", "w") as f:
    f.write(CONTENT)

# ---------------------------------------------------------
# 2. Markdown File
# ---------------------------------------------------------
with open("test-files/test-file.md", "w") as f:
    f.write("# Client Document Portal — Test File\n\n")
    f.write("_This document is auto‑generated for testing purposes only._\n\n")
    f.write("```\n" + CONTENT + "\n```")

# ---------------------------------------------------------
# 3. HTML File (with brand colors)
# ---------------------------------------------------------
HTML = f"""
<!DOCTYPE html>
<html>
<head>
<style>
body {{
    font-family: Arial, sans-serif;
    background: #FAFAFA;
    color: #23272F;
    padding: 40px;
}}
h1 {{ color: #57C4DC; }}
h2 {{ color: #2A2F38; border-bottom: 2px solid #70CDE1; }}
.watermark {{ color: #E5E7EB; font-size: 0.9em; }}
.footer {{ margin-top: 40px; font-size: 0.85em; color: #6B7280; }}
pre {{
    background: #F3F4F6;
    padding: 20px;
    border-radius: 8px;
}}
</style>
</head>
<body>

<h1>Client Document Portal — Test File</h1>
<p class="watermark">[TEST FILE — NOT FOR PRODUCTION]</p>

<pre>{CONTENT}</pre>

<div class="footer">Generated for testing purposes only — williamjwhite.me</div>

</body>
</html>
"""

with open("test-files/test-file.html", "w") as f:
    f.write(HTML)

# ---------------------------------------------------------
# 4. DOCX File
# ---------------------------------------------------------
doc = Document()
doc.add_heading("Client Document Portal — Test File", level=1)
doc.add_paragraph("[TEST FILE — NOT FOR PRODUCTION]")
for line in CONTENT.split("\n"):
    doc.add_paragraph(line)
doc.add_paragraph("Generated for testing purposes only — williamjwhite.me")
doc.save("test-files/test-file.docx")

# ---------------------------------------------------------
# 5. PDF File
# ---------------------------------------------------------
pdf_path = "test-files/test-file.pdf"
c = canvas.Canvas(pdf_path, pagesize=letter)
text = c.beginText(40, 750)
text.setFont("Helvetica", 10)

for line in CONTENT.split("\n"):
    text.textLine(line)

c.drawText(text)
c.setFont("Helvetica-Oblique", 8)
c.drawString(40, 20, "Test File — williamjwhite.me")
c.save()

# ---------------------------------------------------------
# 6. PNG & JPG (simple watermark images)
# ---------------------------------------------------------
with open("test-files/test-file.png", "wb") as f:
    f.write(b"\x89PNG\r\n\x1a\nTEST FILE — NOT FOR PRODUCTION")

with open("test-files/test-file.jpg", "wb") as f:
    f.write(b"\xff\xd8\xffTEST FILE — NOT FOR PRODUCTION")

print("All test files generated in ./test-files")
